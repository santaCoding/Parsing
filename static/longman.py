import requests
from bs4 import BeautifulSoup
from time import sleep
from random import randint
import docx
from docx.shared import RGBColor
from docx.shared import Pt
from selenium import webdriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By

headers = ({'User-Agent': 'Mozilla/5.0 (Windows NT 6.1) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/41.0.2228.0 Safari/537.36'})
doc = docx.Document()
paragraph_format = doc.styles['Normal'].paragraph_format
paragraph_format.line_spacing = Pt(12)
url = 'https://www.ldoceonline.com/browse/english/0-9/'
response = requests.get(url, headers=headers)

def getInnerData(a):
    response = requests.get(a, headers=headers)
    if response.status_code == 200:
        response = requests.get(a, headers=headers)
        html_soup_inner = BeautifulSoup(response.text, 'html.parser')
        title = HYPHENATION = PronCodes = POS = word_def = span = ''
        title = html_soup_inner.find('h1', class_=('pagetitle')).get_text()
        title = HYPHENATION = PronCodes = POS = word_def = span = gram = add = deriv = pos2 = gram2 = ''
        title = html_soup_inner.find('h1', class_=('pagetitle')).get_text()
        try:
            HYPHENATION = html_soup_inner.find('span', class_=('HYPHENATION')).get_text()
        except:
            pass
        try:
            PronCodes = html_soup_inner.find('span', class_=('PronCodes')).get_text()
        except:
            pass
        try:
            POS = html_soup_inner.find('span', class_=('POS')).get_text()
        except:
            pass
        word_def = html_soup_inner.find('span', class_=('DEF')).get_text()
        try:
            span = html_soup_inner.find('span', class_=('dictionary_intro','span')).get_text()
        except:
            pass
        try:
            gram = html_soup_inner.find('span', class_=('GRAM')).get_text()
        except:
            pass
        try:
            deriv = html_soup_inner.find('span', class_=('RunOn','current')).find('span', class_=('DERIV')).get_text()
            pos2 = html_soup_inner.find('span', class_=('RunOn', 'current')).find('span', class_=('POS')).get_text()
            gram2 = html_soup_inner.find('span', class_=('RunOn', 'current')).find('span', class_=('GRAM')).get_text()
        except:
            pass
        t = doc.add_heading(0)
        titles = t.add_run(title)
        titles.font.size = Pt(18)
        titles.font.color.rgb = RGBColor(0, 0, 0)
        titles.font.name = 'Arial'
        if span != '':
            sp = doc.add_heading(level=5)
            sps = sp.add_run(span)
            sps.font.size = Pt(11)
            sps.font.name = "Arial"
        if HYPHENATION != '':
            hyp = doc.add_heading(0)
            hyps = hyp.add_run(HYPHENATION)
            hyps.font.size = Pt(18)
            hyps.font.name = 'Arial'
            hyps.font.color.rgb = RGBColor(255, 0, 0)
        if PronCodes != '':
            pron = doc.add_heading(level=4)
            prons = pron.add_run(PronCodes)
            prons.font.name = "Arial"
            prons.font.color.rgb = RGBColor(0, 0, 0)
        if POS != '':
            pos = doc.add_heading(level=4)
            poses = pos.add_run(POS)
            poses.font.color.rgb = RGBColor(0, 128, 0)
            poses.font.name = 'Arial'
            poses.italic = False
        if gram != '':
            gramy = doc.add_paragraph()
            gramys = gramy.add_run(gram)
            gramys.font.name = 'Arial'
            gramys.font.color.rgb = RGBColor(0, 128, 0)
        if add != '':
            addy = doc.add_paragraph()
            adds = addy.add_run(add)
            adds.font.color.rgb = RGBColor(0, 0, 0)
        if deriv != '':
            derivy = doc.add_paragraph()
            derivs = derivy.add_run(deriv)
            derivs.font.name = 'Arial'
            posy = doc.add_paragraph()
            poses2 = posy.add_run(pos2)
            derivs.font.name = 'Arial'
            gramy2 = doc.add_paragraph()
            gramys2 = gramy2.add_run(gram2)
            gramys2.font.name = 'Arial'
        text = doc.add_paragraph()
        texts = text.add_run(word_def)
        texts.font.name = 'Arial'
        doc.add_paragraph('\n\n')

def getDefs(a):
    response = requests.get(a, headers=headers)
    if response.status_code == 200:
        print('--', a)
        response = requests.get(a, headers=headers)
        html_soup_inner = BeautifulSoup(response.text, 'html.parser')
        try:
            title = HYPHENATION = PronCodes = POS = word_def = span = ''
            title = html_soup_inner.find('h1', class_=('pagetitle')).get_text()
            title = HYPHENATION = PronCodes = POS = word_def = span = gram = add = deriv = pos2 = gram2 = ''
            title = html_soup_inner.find('h1', class_=('pagetitle')).get_text()
            try:
                HYPHENATION = html_soup_inner.find('span', class_=('HYPHENATION')).get_text()
            except:
                pass
            try:
                PronCodes = html_soup_inner.find('span', class_=('PronCodes')).get_text()
            except:
                pass
            try:
                POS = html_soup_inner.find('span', class_=('POS')).get_text()
            except:
                pass
            word_def = html_soup_inner.find('span', class_=('DEF')).get_text()
            try:
                span = html_soup_inner.find('span', class_=('dictionary_intro','span')).get_text()
            except:
                pass
            try:
                gram = html_soup_inner.find('span', class_=('GRAM')).get_text()
            except:
                pass
            try:
                deriv = html_soup_inner.find('span', class_=('RunOn','current')).find('span', class_=('DERIV')).get_text()
                pos2 = html_soup_inner.find('span', class_=('RunOn', 'current')).find('span', class_=('POS')).get_text()
                gram2 = html_soup_inner.find('span', class_=('RunOn', 'current')).find('span', class_=('GRAM')).get_text()
            except:
                pass
            t = doc.add_heading(0)
            titles = t.add_run(title)
            titles.font.size = Pt(18)
            titles.font.color.rgb = RGBColor(0, 0, 0)
            titles.font.name = 'Arial'
            if span != '':
                sp = doc.add_heading(level=5)
                sps = sp.add_run(span)
                sps.font.size = Pt(11)
                sps.font.name = "Arial"
            if HYPHENATION != '':
                hyp = doc.add_heading(0)
                hyps = hyp.add_run(HYPHENATION)
                hyps.font.size = Pt(18)
                hyps.font.name = 'Arial'
                hyps.font.color.rgb = RGBColor(255, 0, 0)
            if PronCodes != '':
                pron = doc.add_heading(level=4)
                prons = pron.add_run(PronCodes)
                prons.font.name = "Arial"
                prons.font.color.rgb = RGBColor(0, 0, 0)
            if POS != '':
                pos = doc.add_heading(level=4)
                poses = pos.add_run(POS)
                poses.font.color.rgb = RGBColor(0, 128, 0)
                poses.font.name = 'Arial'
                poses.italic = False
            if gram != '':
                gramy = doc.add_paragraph()
                gramys = gramy.add_run(gram)
                gramys.font.name = 'Arial'
                gramys.font.color.rgb = RGBColor(0, 128, 0)
            if add != '':
                addy = doc.add_paragraph()
                adds = addy.add_run(add)
                adds.font.color.rgb = RGBColor(0, 0, 0)
            if deriv != '':
                derivy = doc.add_paragraph()
                derivs = derivy.add_run(deriv)
                derivs.font.name = 'Arial'
                posy = doc.add_paragraph()
                poses2 = posy.add_run(pos2)
                poses2.font.name = 'Arial'
                poses2.font.color.rgb = RGBColor(0, 128, 0)
                gramy2 = doc.add_paragraph()
                gramys2 = gramy2.add_run(gram2)
                gramys2.font.name = 'Arial'
                gramys2.font.color.rgb = RGBColor(0, 128, 0)
            text = doc.add_paragraph()
            texts = text.add_run(word_def)
            texts.font.name = 'Arial'
            doc.add_paragraph('\n\n')
        except:
            a_list = ['https://www.ldoceonline.com' + x['href'] for x in html_soup_inner.find('span', class_=('ldoceEntry', 'Entry')).findAll('a')]
            print(a_list)
            words = []
            for a in a_list:
                getInnerData(a)
                break

def getGroups(a, file_count):
    response = requests.get(a, headers=headers)
    if response.status_code == 200:
        html_soup_inner = BeautifulSoup(response.text, 'html.parser')
        a_list = ['https://www.ldoceonline.com' + x['href'] for x in html_soup_inner.find('ul', class_=('browse_results')).findAll('a')]
        for a in a_list:
            getDefs(a)
            



if response.status_code == 200:
    html_soup = BeautifulSoup(response.text, 'html.parser')
    a_list = ['https://www.ldoceonline.com' + x['href'] for x in html_soup.find('ul', class_=('browse_groups')).findAll('a')]
    print(a_list)
    file_count = 0
    for a in a_list:
        file_count += 1
        getGroups(a, file_count)
        
    doc.save('0-9.docx')
    