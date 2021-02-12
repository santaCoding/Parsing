import docx
document = docx.Document()
run = document.add_paragraph().add_run()

style = document.styles['Normal']
font = style.font
font.name = 'Arial Bold'
font.size = docx.shared.Pt(15)
document.add_paragraph('Some text').add_run()
document.save('dwdw.docx')